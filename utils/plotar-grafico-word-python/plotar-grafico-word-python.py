# pip install python-docx

import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from PIL import Image
import io



# primeiro tranformamos os dados em um dataframe
df = pd.DataFrame({'mesAno': ['Jan/22', 'Fev/22','Mar/22','Abr/22', 'Mai/22', 'Jun/22'],
                    'acessos': [2005, 1589, 2123, 2193, 1235, 2410]})

# Configura o gráfico
df.plot(kind='bar', x='mesAno', y='acessos',  title='Total de Acessos Mensal', rot=0)     


# Armazena o gráfico em bytes
figure = io.BytesIO(); 
plt.savefig(figure); 
figure.seek(0); 


# instância o documento
document = Document()

# Adiciona um título 
document.add_heading('Plotando Gráfico com Python', level=0)


# Adiciona uma imagem ao documento, informando o tamanho da imagem em polegadas
document.add_picture(figure, width=Inches(5.25))


# Adiciona uma quebra de linha ao documento
document.add_page_break()

# Salva o documento gerado em um arquivo
document.save('documento.docx')