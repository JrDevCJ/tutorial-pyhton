# pip install python-docx
from docx import Document
from docx.shared import Inches

# instância o documento
document = Document()

# Adiciona um título 
document.add_heading('Título do Documento', level=0)

# Adiciona um paragrafo e configura a estilização
p = document.add_paragraph('Um simples parágrafo, ')
p.add_run('frase com estilo bold').bold = True  # Estiliza a frase com bold
p.add_run('frase com estilo italic.').italic = True #Estiliza a frase com italic

# Adiciona um subtítulo ao documento 
document.add_heading('SubTítulo do Documento, level 1', level=1)

# Adiciona um parágrafo como uma lista
document.add_paragraph(
    'Item da lista', style='List Bullet'
)

# Adiciona um parágrafo como uma lista ordenada
document.add_paragraph(
    'Item ordenado', style='List Number'
)

# Adiciona uma imagem ao documento, informando o tamanho da imagem em polegadas
document.add_picture('./python-logo.png', width=Inches(1.25))

# Cria uma tupla coms os dados a serem inseridos na tabela
dados = (
    (1, 'Carla'),
    (2, 'Paulo'),
    (3, 'Lara')
)

# Cria a tabela, com 1 linha e 2 colunas
table = document.add_table(rows=1, cols=2)

# Armazena a primeira linha da tabela, para popular o cabeçalho 
hdr_cells = table.rows[0].cells
hdr_cells[0].text =  'Id'
hdr_cells[1].text = 'Nome'

# Percorre a tupla e preenche as informações na tabela
for id, nome in dados:
    row_cells = table.add_row().cells
    row_cells[0].text = str(id)
    row_cells[1].text = nome

# Adiciona uma quebra de linha ao documento
document.add_page_break()

# Salva o documento gerado em um arquivo
document.save('documento.docx')